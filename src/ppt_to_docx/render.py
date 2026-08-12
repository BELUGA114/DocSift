from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document

from .models import Organization, RunPaths


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    escape = lambda value: value.replace("|", "\\|").replace("\n", "<br>")
    header = "| " + " | ".join(escape(value) for value in normalized[0]) + " |"
    divider = "| " + " | ".join("---" for _ in range(columns)) + " |"
    body = ["| " + " | ".join(escape(value) for value in row) + " |" for row in normalized[1:]]
    return "\n".join([header, divider, *body])


def _render_markdown(organization: Organization, index: list[dict[str, object]]) -> str:
    lines = [f"# {organization.title}", ""]
    for unit in organization.units:
        if unit.heading:
            lines.extend([f"{'#' * min(max(unit.level + 1, 2), 6)} {unit.heading}", ""])
        for paragraph in unit.paragraphs:
            lines.extend([paragraph, ""])
        for table_data in unit.tables:
            table = _markdown_table(table_data.rows)
            if table:
                lines.extend([table, ""])
    lines.extend(["## 待确认内容", ""])
    for item in organization.uncertain_items:
        lines.append(f"- {item.text}（来源：{', '.join(item.sources)}；原因：{item.reason}）")
    lines.extend(["", "## 来源索引", ""])
    for item in index:
        lines.append(f"- {item['heading'] or '无标题内容'}：{', '.join(item['sources'])}")
    return "\n".join(lines).rstrip() + "\n"


def render_document(paths: RunPaths, organization: Organization) -> Path:
    document = Document()
    document.add_heading(organization.title, 0)
    index: list[dict[str, object]] = []
    for unit in organization.units:
        if unit.heading:
            document.add_heading(unit.heading, min(max(unit.level, 1), 9))
        for paragraph in unit.paragraphs:
            document.add_paragraph(paragraph)
        for table_data in unit.tables:
            if not table_data.rows:
                continue
            columns = max(len(row) for row in table_data.rows)
            table = document.add_table(rows=0, cols=columns)
            table.style = "Table Grid"
            for row_data in table_data.rows:
                cells = table.add_row().cells
                for index_value, value in enumerate(row_data):
                    cells[index_value].text = value
        index.append({"heading": unit.heading, "sources": unit.sources})
    document.add_heading("待确认内容", 1)
    for item in organization.uncertain_items:
        document.add_paragraph(f"{item.text}（来源：{', '.join(item.sources)}；原因：{item.reason}）")
    document.add_heading("来源索引", 1)
    for item in index:
        document.add_paragraph(f"{item['heading'] or '无标题内容'}：{', '.join(item['sources'])}")
    paths.output_dir.mkdir(parents=True, exist_ok=True)
    output_format = os.environ.get("OUTPUT_FORMAT", "docx").lower()
    if output_format not in {"docx", "markdown", "md"}:
        raise ValueError("OUTPUT_FORMAT 必须是 docx 或 markdown")
    configured_name = os.environ.get("OUTPUT_NAME") or os.environ.get("OUTPUT_DOCUMENT_NAME")
    output_stem = Path(configured_name or "整理结果").stem
    if output_format in {"markdown", "md"}:
        output_path = paths.output_dir / f"{output_stem}.md"
        output_path.write_text(_render_markdown(organization, index), encoding="utf-8")
    else:
        output_path = paths.output_dir / f"{output_stem}.docx"
        document.save(output_path)
    (paths.output_dir / "source-index.json").write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    return output_path
